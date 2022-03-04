import docx
from typing import Dict
from collections import OrderedDict
import re


def get_Q_para_num(doc: docx.Document) -> Dict:
    qp_dict = OrderedDict()
    for ip, para in enumerate(doc.paragraphs):
        qnum = re.findall('Q(\d+)', para.text)
        # if para.text.startswith('Q')
        if qnum:
            if len(qnum) > 1:
                print(f'Found many Qs: {qnum}, skipping...')
                continue
            curr_question = qnum[0]
            qp_dict[curr_question] = {}
            qp_dict[curr_question]['para_num'] = ip
    return qp_dict


def find_next_para_nums(q_para_dict: Dict) -> Dict:
    for i in range(len(q_para_dict)):
        q, _ = list(q_para_dict.items())[i]
        if i < len(q_para_dict) - 1:
            q_para_dict[q]['para_num_next'] = list(q_para_dict.items())[i+1][1]['para_num']
        else:
            q_para_dict['para_num_next'] = float('nan')
    return q_para_dict


def get_questions(doc: docx.Document, q_para_dict: Dict) -> Dict:
    for q, qd in q_para_dict.items():
        q_para = doc.paragraphs[qd['para_num']]
        billy = 1
        req = re.match('Q\d+', q_para.text)
        if req.groups():
            s_after_q = req.span()[1]
    pass


def get_answers(doc: docx.Document, q_para_dict: Dict) -> Dict:
    for q, qd in q_para_dict.items():
        for para in doc.paragraphs[qd['para_num'] + 1 : qd['para_num_next'] - 1]:
            billy = 2
            # TODO: walk through answers
    pass


def run_parser(docx_filepath: str) -> None:
    doc = docx.Document(docx_filepath)
    q_para_dict = get_Q_para_num(doc)
    q_para_dict = find_next_para_nums(q_para_dict)
    q_para_dict = get_questions(doc, q_para_dict)
    q_para_dict = get_answers(doc, q_para_dict)
    pass


if __name__ == '__main__':
    docx_filepath = '/Volumes/Google Drive/My Drive/Projects/Med_Sadhvi/Female Sexual Dysfunction/WISH Survey (1).docx'
    run_parser(docx_filepath)
